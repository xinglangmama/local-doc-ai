"""文档处理模块
包含文档文本提取、文本分割、向量存储等功能"""

import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.schema import Document
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from local_embeddings import LocalEmbeddings
import os

# 全局变量，用于存储处理后的文档块
processed_chunks = []

def get_embeddings():
    """
    获取嵌入模型实例，使用Streamlit session_state避免重复初始化
    """
    if "embeddings_instance" not in st.session_state:
        st.session_state.embeddings_instance = LocalEmbeddings()
    return st.session_state.embeddings_instance


def get_pdf_text(pdf_docs):
    """
    从上传的PDF文件列表中提取所有文本内容
    参数: pdf_docs - PDF文件列表
    返回: 合并后的文本字符串
    """
    text = ""  # 初始化空字符串用于存储提取的文本
    # 遍历每个PDF文件
    for pdf in pdf_docs:
        try:
            pdf_reader = PdfReader(pdf)  # 创建PDF阅读器对象
            # 遍历PDF的每一页
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()  # 提取页面文本
                    if page_text:
                        text += page_text + "\n"  # 添加换行符分隔页面
                except Exception as e:
                    st.warning(f"⚠️ PDF第{page_num+1}页文本提取失败: {str(e)}")
                    continue
            
            if not text.strip():
                st.warning(f"⚠️ PDF文件 {pdf.name} 中未提取到任何文本内容")
        except Exception as e:
            st.error(f"❌ 读取PDF文件 {pdf.name} 时发生错误: {str(e)}")
            continue
    
    return text  # 返回合并后的所有文本


def get_docx_text(docx_docs):
    """
    从上传的Word文档列表中提取所有文本内容
    参数: docx_docs - Word文档文件列表
    返回: 合并后的文本字符串
    """
    text = ""  # 初始化空字符串用于存储提取的文本
    # 遍历每个Word文档文件
    for docx_file in docx_docs:
        try:
            doc = docx.Document(docx_file)  # 创建Word文档对象
            doc_text = ""
            # 遍历文档中的每个段落
            for para in doc.paragraphs:
                if para.text.strip():  # 只添加非空段落
                    doc_text += para.text + "\n"  # 提取段落文本并添加换行符
            
            if doc_text.strip():
                text += doc_text
            else:
                st.warning(f"⚠️ Word文档 {docx_file.name} 中未提取到任何文本内容")
                
        except Exception as e:
            st.error(f"❌ 读取Word文档 {docx_file.name} 时发生错误: {str(e)}")
            continue
    
    return text  # 返回合并后的所有文本


def get_text_chunks(text):
    """
    将长文本分割成较小的文本块，便于向量化和检索
    参数: text - 需要分割的原始文本字符串
    返回: 文本块列表
    """
    if not text or not text.strip():
        st.warning("⚠️ 输入文本为空，无法进行分块处理")
        return []
    
    try:
        # 创建递归字符文本分割器，用于智能分割文本
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=300,        # 每个文本块的最大字符数
            chunk_overlap=50,      # 相邻文本块之间的重叠字符数，保持上下文连贯性
            length_function=len     # 使用字符长度作为分割依据
        )
        # 使用分割器将文本分割成多个块
        chunks = text_splitter.split_text(text)
        
        # 过滤掉空的或过短的文本块
        valid_chunks = [chunk.strip() for chunk in chunks if chunk.strip() and len(chunk.strip()) > 10]
        
        if not valid_chunks:
            st.warning("⚠️ 文本分块后没有有效的文本块")
            return []
        
        return valid_chunks  # 返回分割后的有效文本块列表
    except Exception as e:
        st.error(f"❌ 文本分块过程中发生错误: {str(e)}")
        return []


def get_vectorstore(text_chunks, source_filename="未知文档", append_mode=True):
    """
    将文本块转换为向量并创建FAISS向量数据库
    参数: 
        text_chunks - 文本块列表
        source_filename - 源文档文件名
        append_mode - 是否追加模式（True: 追加到现有数据库，False: 创建新数据库）
    返回: FAISS向量存储对象
    """
    if not text_chunks:
        st.warning("⚠️ 文本块列表为空，无法创建向量数据库")
        return None
    
    # 过滤掉空的文本块
    valid_chunks = [chunk.strip() for chunk in text_chunks if chunk and chunk.strip()]
    
    if not valid_chunks:
        st.warning("⚠️ 没有有效的文本块，无法创建向量数据库")
        return None
    
    try:
        # 获取共享的嵌入模型实例，避免重复初始化
        st.info("📡 正在获取嵌入模型...")
        embeddings = get_embeddings()
        
        # 显示文本块样本（用于调试）
        st.info(f"📝 文本块样本: {valid_chunks[0][:100]}...")
        
        # 创建Document对象列表，为每个文本块添加metadata
        st.info("📄 正在创建Document对象并添加metadata...")
        documents = [
            Document(page_content=chunk, metadata={"source": source_filename, "chunk_id": i})
            for i, chunk in enumerate(valid_chunks)
        ]
        
        # 检查是否存在现有的向量数据库且启用追加模式
        if append_mode and check_vectorstore_exists():
            st.info("🔄 检测到现有向量数据库，正在追加新内容...")
            try:
                # 加载现有的向量数据库
                existing_vectorstore = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
                
                # 创建新文档的向量数据库
                new_vectorstore = FAISS.from_documents(documents, embeddings)
                
                # 将新向量数据库合并到现有数据库中
                existing_vectorstore.merge_from(new_vectorstore)
                vectorstore = existing_vectorstore
                
                st.success(f"✅ 成功追加 {len(valid_chunks)} 个新文本块到现有向量数据库！")
            except Exception as e:
                st.warning(f"⚠️ 加载现有数据库失败，将创建新数据库: {str(e)}")
                # 如果加载失败，创建新的向量数据库
                st.info("🔄 正在创建新的向量数据库...")
                vectorstore = FAISS.from_documents(documents, embeddings)
                st.success(f"✅ 新向量数据库创建成功！共处理了 {len(valid_chunks)} 个文本块。")
        else:
            # 创建新的向量数据库
            st.info(f"🔄 正在创建新的向量数据库，处理 {len(valid_chunks)} 个文本块...")
            vectorstore = FAISS.from_documents(documents, embeddings)
            st.success(f"✅ 新向量数据库创建成功！共处理了 {len(valid_chunks)} 个文本块。")
        
        # 保存向量数据库到本地文件，实现持久化存储
        st.info("💾 正在保存向量数据库到本地...")
        vectorstore.save_local("faiss_index")
        st.success("✅ 向量数据库已保存到本地！")
        
        return vectorstore  # 返回创建的向量存储对象
    except Exception as e:
        st.error(f"❌ 创建向量数据库时发生错误：{str(e)}")
        st.error(f"错误详情: {type(e).__name__}")
        import traceback
        st.error(f"完整错误信息: {traceback.format_exc()}")
        return None


def load_vectorstore():
    """
    从本地文件加载已保存的FAISS向量数据库
    返回: FAISS向量存储对象，如果文件不存在则返回None
    """
    try:
        # 获取共享的嵌入模型实例（必须与保存时使用的模型一致）
        embeddings = get_embeddings()
        # 从本地文件加载向量数据库
        vectorstore = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
        return vectorstore
    except Exception as e:
        # 如果加载失败（如文件不存在），返回None
        return None


def process_uploaded_files(uploaded_files, append_mode=False):
    """
    处理上传的文件，提取文本并创建向量数据库
    
    Args:
        uploaded_files: Streamlit上传的文件列表
        append_mode: 是否为追加模式，True表示追加到现有数据库，False表示创建新数据库
        
    Returns:
        FAISS向量数据库对象，如果处理失败则返回None
    """
    if not uploaded_files:
        st.warning("📋 请先上传文档！")
        return None
    
    # 显示处理进度指示器
    with st.spinner("🔄 正在处理文档，请稍候..."):
        raw_text = ""  # 初始化变量存储所有提取的文本
        file_names = []  # 存储文件名列表
        
        # 从上传文件中筛选出PDF文件
        pdf_files = [f for f in uploaded_files if f.type == "application/pdf"]
        # 从上传文件中筛选出Word文档文件
        docx_files = [f for f in uploaded_files if f.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]

        # 如果有PDF文件，提取其中的文本内容
        if pdf_files:
            raw_text += get_pdf_text(pdf_files)
            file_names.extend([f.name for f in pdf_files])
        # 如果有Word文档，提取其中的文本内容
        if docx_files:
            raw_text += get_docx_text(docx_files)
            file_names.extend([f.name for f in docx_files])

        # 检查是否成功提取到文本内容
        if not raw_text.strip():
            st.error("❌ 未能从上传的文档中提取到有效文本内容！请检查文档格式或内容。")
            return None
        
        # 显示提取到的文本统计信息
        st.info(f"📄 成功提取文本，总字符数: {len(raw_text)}")
        
        # 合并所有文件名作为源文档标识
        source_filename = ", ".join(file_names) if file_names else "未知文档"

        # 将提取的原始文本分割成小块，便于向量化处理
        text_chunks = get_text_chunks(raw_text)
        
        # 验证文本分块结果
        if not text_chunks:
            st.error("❌ 文本分块失败，无法创建向量数据库！")
            return None
        
        st.info(f"📝 文本已分割为 {len(text_chunks)} 个块")

        # 将文本块转换为向量并创建可搜索的向量数据库
        vectorstore = get_vectorstore(text_chunks, source_filename, append_mode)
        
        if vectorstore is None:
            st.error("❌ 文档处理失败，请检查上传的文件！")
            return None
        
        return vectorstore


def check_vectorstore_exists():
    """
    检查本地向量数据库是否存在
    返回: bool - 数据库是否存在
    """
    return os.path.exists("faiss_index")


def get_vectorstore_info():
    """
    获取向量数据库的详细信息
    
    Returns:
        dict: 包含chunk个数和文档来源列表的字典，如果数据库不存在则返回None
    """
    if not check_vectorstore_exists():
        return None
    
    try:
        # 加载向量数据库
        embeddings = LocalEmbeddings()
        vectorstore = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
        
        # 获取所有文档的metadata
        docstore = vectorstore.docstore
        all_docs = [docstore.search(doc_id) for doc_id in docstore._dict.keys()]
        
        # 统计chunk个数
        chunk_count = len(all_docs)
        
        # 提取文档来源列表
        source_files = set()
        for doc in all_docs:
            if hasattr(doc, 'metadata') and 'source' in doc.metadata:
                source_files.add(doc.metadata['source'])
        
        return {
            'chunk_count': chunk_count,
            'source_files': list(source_files)
        }
    
    except Exception as e:
        print(f"获取向量数据库信息时出错: {e}")
        return None